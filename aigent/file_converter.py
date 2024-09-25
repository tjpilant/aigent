# File: file_converter.py
# Author: Tj Pilant
# Description: Handles conversion of various file types to JSONL and other formats
# Version: 0.5.2

import json
import logging
import os
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime

import pdfplumber
import PyPDF2
from docx import Document

from aigent.image_converter import ImageConverter
from aigent.models import AgentTraits, ProjectInfo

# Set up logging
logging.basicConfig(
    filename="file_converter.log",
    level=logging.DEBUG,
    format="%(asctime)s:%(levelname)s:%(message)s",
)

class FileConverter:
    def __init__(self):
        self.image_converter = ImageConverter()

    def convert_file(self, input_file, output_dir, project_info, agent_traits, output_formats=['jsonl'], use_ocr=True, use_cloud_vision=False):
        logging.debug("Starting convert_file method")
        logging.debug(f"Type of project_info: {type(project_info)}")
        logging.debug(f"Type of agent_traits: {type(agent_traits)}")
        
        file_extension = os.path.splitext(input_file)[1].lower()
        
        if file_extension == '.pdf':
            return self.convert_pdf(input_file, output_dir, project_info, agent_traits, output_formats, use_ocr, use_cloud_vision)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.tif']:
            return self.convert_image(input_file, output_dir, project_info, agent_traits, output_formats, use_ocr, use_cloud_vision)
        elif file_extension == '.docx':
            return self.convert_docx(input_file, output_dir, project_info, agent_traits, output_formats)
        elif file_extension == '.txt':
            return self.convert_txt(input_file, output_dir, project_info, agent_traits, output_formats)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def convert_pdf(self, input_file, output_dir, project_info, agent_traits, output_formats, use_ocr, use_cloud_vision):
        try:
            if use_ocr:
                pages = []
                with pdfplumber.open(input_file) as pdf:
                    for page in pdf.pages:
                        img = page.to_image()
                        img_path = f"temp_page_{page.page_number}.png"
                        img.save(img_path)
                        ocr_text = self.image_converter.perform_ocr(img_path, use_cloud_vision)
                        pages.append(ocr_text)
                        os.remove(img_path)
            else:
                with pdfplumber.open(input_file) as pdf:
                    pages = [page.extract_text() for page in pdf.pages]
            
            return self.process_text(pages, input_file, output_dir, project_info, agent_traits, output_formats)
        except Exception as e:
            logging.error(f"Error converting PDF: {str(e)}")
            raise

    def convert_image(self, input_file, output_dir, project_info, agent_traits, output_formats, use_ocr, use_cloud_vision):
        try:
            if use_ocr:
                text = self.image_converter.perform_ocr(input_file, use_cloud_vision)
            else:
                raise ValueError("OCR must be enabled for image conversion")
            
            pages = [text]  # Treat the entire OCR result as a single page
            
            return self.process_text(pages, input_file, output_dir, project_info, agent_traits, output_formats)
        except Exception as e:
            logging.error(f"Error converting image: {str(e)}")
            raise

    def convert_docx(self, input_file, output_dir, project_info, agent_traits, output_formats):
        try:
            doc = Document(input_file)
            pages = [paragraph.text for paragraph in doc.paragraphs]
            
            return self.process_text(pages, input_file, output_dir, project_info, agent_traits, output_formats)
        except Exception as e:
            logging.error(f"Error converting DOCX: {str(e)}")
            raise

    def convert_txt(self, input_file, output_dir, project_info, agent_traits, output_formats):
        try:
            with open(input_file, 'r', encoding='utf-8') as file:
                text = file.read()
            pages = [text]  # Treat the entire text file as a single page
            
            return self.process_text(pages, input_file, output_dir, project_info, agent_traits, output_formats)
        except Exception as e:
            logging.error(f"Error converting TXT: {str(e)}")
            raise

    def process_text(self, pages, input_file, output_dir, project_info, agent_traits, output_formats):
        output_files = {}
        
        # Always generate JSONL
        jsonl_output = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(input_file))[0]}.jsonl")
        self.save_as_jsonl(pages, jsonl_output, project_info, agent_traits)
        output_files['jsonl'] = jsonl_output
        
        # Generate additional formats if requested
        if 'txt' in output_formats:
            txt_output = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(input_file))[0]}.txt")
            self.save_as_txt(pages, txt_output)
            output_files['txt'] = txt_output
        
        if 'md' in output_formats:
            md_output = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(input_file))[0]}.md")
            self.save_as_md(pages, md_output)
            output_files['md'] = md_output
        
        if 'docx' in output_formats:
            docx_output = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(input_file))[0]}_converted.docx")
            self.save_as_docx(pages, docx_output)
            output_files['docx'] = docx_output
        
        return output_files

    def save_as_jsonl(self, pages, output_file, project_info, agent_traits):
        logging.debug("Starting save_as_jsonl method")
        logging.debug(f"Type of project_info: {type(project_info)}")
        logging.debug(f"Type of agent_traits: {type(agent_traits)}")
        
        try:
            with open(output_file, 'w', encoding='utf-8') as jsonl_file:
                for i, page_text in enumerate(pages, 1):
                    cleaned_text = self.clean_text(page_text)
                    try:
                        project_info_dict = project_info.to_dict()
                    except AttributeError:
                        logging.error(f"project_info does not have to_dict method. Type: {type(project_info)}")
                        project_info_dict = vars(project_info) if hasattr(project_info, '__dict__') else str(project_info)

                    try:
                        agent_traits_dict = agent_traits.to_dict()
                    except AttributeError:
                        logging.error(f"agent_traits does not have to_dict method. Type: {type(agent_traits)}")
                        agent_traits_dict = vars(agent_traits) if hasattr(agent_traits, '__dict__') else str(agent_traits)

                    entry = {
                        "page_number": i,
                        "content": cleaned_text,
                        "metadata": {
                            "project_info": project_info_dict,
                            "agent_traits": agent_traits_dict
                        }
                    }
                    logging.debug(f"Type of entry: {type(entry)}")
                    logging.debug(f"Type of metadata: {type(entry['metadata'])}")
                    json.dump(entry, jsonl_file)
                    jsonl_file.write('\n')
            
            logging.debug("Completed save_as_jsonl method")
        except Exception as e:
            logging.error(f"Error in save_as_jsonl method: {str(e)}")
            raise

    def clean_text(self, text):
        # Remove line breaks within words
        cleaned_text = re.sub(r'(?<=[a-zA-Z])-\n(?=[a-zA-Z])', '', text)
        # Join words split across lines, but keep paragraph breaks
        cleaned_text = re.sub(r'(?<=[a-zA-Z])\n(?=[a-z])', ' ', cleaned_text)
        # Remove extra whitespace
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
        return cleaned_text

    def save_as_txt(self, pages, output_file):
        with open(output_file, 'w', encoding='utf-8') as txt_file:
            for page in pages:
                txt_file.write(page + '\n\n')

    def save_as_md(self, pages, output_file):
        with open(output_file, 'w', encoding='utf-8') as md_file:
            for i, page in enumerate(pages, 1):
                md_file.write(f"## Page {i}\n\n{page}\n\n")

    def save_as_docx(self, pages, output_file):
        doc = Document()
        for page in pages:
            doc.add_paragraph(page)
        doc.save(output_file)

    def batch_convert_files(self, input_files, output_dir, project_info, agent_traits, output_formats=["jsonl"], use_ocr=True, use_cloud_vision=False):
        logging.info(f"Starting batch_convert_files with {len(input_files)} files")
        logging.debug(f"Type of project_info: {type(project_info)}")
        logging.debug(f"Type of agent_traits: {type(agent_traits)}")
        logging.debug(f"Output formats: {output_formats}")
        logging.debug(f"Use OCR: {use_ocr}")
        logging.debug(f"Use Cloud Vision: {use_cloud_vision}")

        if len(input_files) > 10:
            raise ValueError("Batch processing is limited to a maximum of 10 files.")

        results = {}
        errors = {}

        with ThreadPoolExecutor() as executor:
            future_to_file = {executor.submit(self.convert_file, input_file, output_dir, project_info, agent_traits, output_formats, use_ocr, use_cloud_vision): input_file for input_file in input_files}
            
            for future in as_completed(future_to_file):
                input_file = future_to_file[future]
                try:
                    result = future.result()
                    results[input_file] = result
                    logging.info(f"Successfully processed file: {input_file}")
                except Exception as e:
                    errors[input_file] = str(e)
                    logging.error(f"Error processing file {input_file}: {str(e)}", exc_info=True)

        logging.info(f"Completed batch_convert_files. Processed {len(results)} files, encountered {len(errors)} errors.")
        return results, errors